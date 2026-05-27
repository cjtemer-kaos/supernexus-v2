#!/usr/bin/env python3
import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

class DocumentGeneratorSkill:
    """
    Skill para generar reportes profesionales en formato Word (.docx) y PDF.
    Diseñado para la gema Synthesizer del ecosistema Nexus.
    """
    
    def __init__(self):
        self.name = "document_generator"
        from pathlib import Path
        _project = Path(__file__).resolve().parents[2]
        self.output_dir = str(_project / "output" / "reports")
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def generate_docx(self, title: str, content: str, filename: str = "report.docx") -> dict:
        """Genera un archivo Word (.docx) con el título y contenido proporcionado."""
        if not filename.endswith(".docx"):
            filename += ".docx"
        
        path = os.path.join(self.output_dir, filename)
        try:
            doc = Document()
            doc.add_heading(title, 0)
            
            # Dividir el contenido por párrafos
            paragraphs = content.split('\n')
            for p in paragraphs:
                if p.strip():
                    doc.add_paragraph(p.strip())
            
            doc.save(path)
            return {"status": "success", "message": f"Archivo Word generado en {path}", "path": path}
        except Exception as e:
            return {"status": "error", "message": f"Error generando Word: {str(e)}"}

    def generate_pdf(self, title: str, content: str, filename: str = "report.pdf") -> dict:
        """Genera un archivo PDF profesional con el título y contenido proporcionado."""
        if not filename.endswith(".pdf"):
            filename += ".pdf"
            
        path = os.path.join(self.output_dir, filename)
        try:
            doc = SimpleDocTemplate(path, pagesize=letter)
            styles = getSampleStyleSheet()
            elements = []
            
            # Título
            elements.append(Paragraph(title, styles['Title']))
            elements.append(Spacer(1, 12))
            
            # Contenido
            paragraphs = content.split('\n')
            for p in paragraphs:
                if p.strip():
                    elements.append(Paragraph(p.strip(), styles['Normal']))
                    elements.append(Spacer(1, 12))
            
            doc.build(elements)
            return {"status": "success", "message": f"Archivo PDF generado en {path}", "path": path}
        except Exception as e:
            return {"status": "error", "message": f"Error generando PDF: {str(e)}"}

    def generate_table_report(self, title: str, headers: list, rows: list, filename: str = "table_report.pdf") -> dict:
        """Genera un PDF con una tabla formateada."""
        if not filename.endswith(".pdf"):
            filename += ".pdf"
            
        path = os.path.join(self.output_dir, filename)
        try:
            doc = SimpleDocTemplate(path, pagesize=letter)
            styles = getSampleStyleSheet()
            elements = []
            
            elements.append(Paragraph(title, styles['Title']))
            elements.append(Spacer(1, 12))
            
            # Crear Tabla
            data = [headers] + rows
            t = Table(data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(t)
            doc.build(elements)
            return {"status": "success", "message": f"Reporte de tabla generado en {path}", "path": path}
        except Exception as e:
            return {"status": "error", "message": f"Error generando tabla PDF: {str(e)}"}

    def info(self):
        return {
            "name": self.name,
            "description": "Generador de documentos Office y PDF para el ecosistema Nexus.",
            "capabilities": ["docx", "pdf", "tables"],
            "output_directory": self.output_dir
        }
